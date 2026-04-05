import pytest
from pathlib import Path
from docx import Document
from replacer import replace_in_docx


def make_docx(tmp_path, paragraphs: list[str]) -> str:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    file_path = tmp_path / "test.docx"
    doc.save(file_path)
    return str(file_path)


def test_basic_english_replacement(tmp_path):
    path = make_docx(tmp_path, ["Hello world"])
    output, count = replace_in_docx(path, "world", "Python")
    assert count == 1
    doc = Document(output)
    assert doc.paragraphs[0].text == "Hello Python"


def test_chinese_replacement(tmp_path):
    path = make_docx(tmp_path, ["我的書包"])
    output, count = replace_in_docx(path, "的", "之")
    assert count == 1
    doc = Document(output)
    assert doc.paragraphs[0].text == "我之書包"


def test_no_match(tmp_path):
    path = make_docx(tmp_path, ["Hello world"])
    output, count = replace_in_docx(path, "xyz", "abc")
    assert count == 0
    doc = Document(output)
    assert doc.paragraphs[0].text == "Hello world"


def test_multiple_occurrences(tmp_path):
    path = make_docx(tmp_path, ["aaa bbb aaa ccc aaa"])
    output, count = replace_in_docx(path, "aaa", "zzz")
    assert count == 3
    doc = Document(output)
    assert doc.paragraphs[0].text == "zzz bbb zzz ccc zzz"


def test_file_not_found():
    with pytest.raises(FileNotFoundError):
        replace_in_docx("nonexistent.docx", "a", "b")


def test_wrong_extension(tmp_path):
    bad_file = tmp_path / "file.txt"
    bad_file.write_text("hello")
    with pytest.raises(ValueError):
        replace_in_docx(str(bad_file), "a", "b")


def test_output_filename(tmp_path):
    path = make_docx(tmp_path, ["hello"])
    output, _ = replace_in_docx(path, "hello", "hi")
    assert Path(output).name == "test_modified.docx"


def test_empty_document(tmp_path):
    path = make_docx(tmp_path, [])
    output, count = replace_in_docx(path, "anything", "else")
    assert count == 0
