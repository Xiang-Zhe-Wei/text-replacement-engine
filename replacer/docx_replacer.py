import re
from pathlib import Path
from docx import Document


def replace_multiple_in_docx(input_path: str, replacements: list[tuple[str, str]]) -> tuple[str, dict]:
    """
    Apply multiple replacements to a .docx file in a single regex pass.
    Longer targets are matched first to prevent shorter rules from interfering
    with the output of more specific rules (e.g. 及→與 overwriting A、B及C).
    Returns (output_path, counts) where counts is a dict of target -> count.
    """
    path = Path(input_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {input_path}")
    if path.suffix.lower() != ".docx":
        raise ValueError(f"Not a .docx file: {input_path}")

    doc = Document(input_path)
    counts = {target: 0 for target, _ in replacements}

    # Sort longest target first so specific rules win over general ones
    sorted_replacements = sorted(replacements, key=lambda x: len(x[0]), reverse=True)
    replacement_map = {t: r for t, r in sorted_replacements}
    pattern = re.compile("|".join(re.escape(t) for t, _ in sorted_replacements))

    def apply(text: str) -> str:
        def replace_match(m: re.Match) -> str:
            matched = m.group(0)
            counts[matched] += 1
            return replacement_map[matched]
        return pattern.sub(replace_match, text)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = apply(run.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = apply(run.text)

    output_path = path.parent / f"{path.stem}_modified{path.suffix}"
    doc.save(output_path)

    return str(output_path), counts


def replace_in_docx(input_path: str, target: str, replacement: str) -> str:
    """
    Replace all occurrences of target with replacement in a .docx file.
    Returns the path of the newly created output file.
    """
    path = Path(input_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {input_path}")
    if path.suffix.lower() != ".docx":
        raise ValueError(f"Not a .docx file: {input_path}")

    doc = Document(input_path)
    count = 0

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if target in run.text:
                count += run.text.count(target)
                run.text = run.text.replace(target, replacement)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if target in run.text:
                            count += run.text.count(target)
                            run.text = run.text.replace(target, replacement)

    output_path = path.parent / f"{path.stem}_modified{path.suffix}"
    doc.save(output_path)

    return str(output_path), count
